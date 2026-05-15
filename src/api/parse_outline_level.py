import re
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.section import CT_SectPr
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table



def getOutlineLevel(inputXml):
    """
    功能从xml字段中提取出<w:outlineLevel w:val="number"/>中的数字number
    参数inputXml
    返回 number
    """
    start_index = inputXml.find('<w:outlineLvl')
    end_index = inputXml.find('>', start_index)
    number = inputXml[start_index:end_index + 1]
    number = re.search(r"\d+", number).group()
    return number


def isTitle(paragraph):
    """
    功能 判断改段落是否设置了大纲等级
    参数paragraph：段落
    返回None:普通正文，没有大纲级别 0:一级标题 2:二级标题 2:三级标题
    """
    # 如果是空行，直接返回None
    if paragraph.text.strip() == '':
        return None

    # 如果该段落是直接在段落里设置大纲级别的，根据xml判断大纲级别
    paragraphXml = paragraph._p.xml
    # print(paragraphXml)
    if paragraphXml.find('<w:outlineLevel') >=0:
        return getOutlineLevel(paragraphXml)
    # 如果该段落是通过样式设置大纲级别的，逐级检索样式及其父样式，判断大纲级别
    targetStyle = paragraphXml.style
    while targetStyle is not None:
        # 如果在该级别style中找到了大纲级别,返回
        # print(targetStyle.element.xml)
        if targetStyle.element.xml.find('<w:outlineLvl') >=0:
            return getOutlineLevel(targetStyle.element.xml)
        else:
            targetStyle = targetStyle.base_style
    # 如果在段落、样式里都没有找到大纲级别，返回None
    return None



def  parse_outline_level(docx_path):
    document = docx.Document(docx_path)
    # 逐段读取docx文档的内容
    result = []
    for paragraph in document.paragraphs:
        level = isTitle(paragraph)
        if level != None and int(level) < 8:
            result.append({"title": paragraph.text.strip(), "level": int(level) +1})
    return result


def parse_docx_level_entry(path):
    outline_level = parse_outline_level(path)
    file = docx.Document(path)
    bodys = file.element.body
    content_str = ""
    result_lst = []
    start_flag = False
    title = ""
    last_level = -1
    begin_text = ""
    for i, child in enumerate(bodys, 1):
        if child.tag.endswith('p'):  # 段落 `/w:document/w:body/w:p`
            para = Paragraph(child, CT_SectPr)
            if not para.text.strip(): continue
            # 段落层级，从"0"开始
            text = para.text.strip()
        elif child.tag.endswith('tbl'):  # 表格`/w:document/w:body/w:tbl`
            tbl = Table(child, CT_tbl)
            table_matrix = []
            for row in tbl.rows:
                tmp_table_row = []
                for cell in row.cells:
                    tmp_table_row.append(str(cell.text))
                table_matrix.append(tmp_table_row)
            text = str(table_matrix)

        if len(outline_level) > 0 and text == outline_level[0]["title"]:
            if start_flag == False:
                start_flag = True
            else:
                result_lst.append({"text": title, "level": last_level, "content": content_str})
            title = text
            content_str = ""
            last_level = outline_level[0]["level"]
            outline_level = outline_level[1:]
        else:
            if start_flag == True:
                content_str += (text + "\n")
            else:
                begin_text += (text + "\n")

    # 最后一个片段
    result_lst.append({"text": title, "level": last_level, "content": content_str})
    # begin_text添加一个大纲等级（有些文档会设置第一个标题的大纲等级）
    if len(result_lst) > 0:
        begin_text = begin_text + "\n" + result_lst[0]["text"] + "\n" + result_lst[0]["content"]

    return result_lst, begin_text

if __name__ == '__main__':
    parse_outline_level(r"C:\Users\李灵佳\Desktop\1.docx")


#if __name__ == '__main__':
#    filepath = r"C:\Users\总体技术中心\Desktop\0-RCL671B地面雷达干扰系统性能鉴定试验总体方案1.1版-20221011-灰度.docx"

#    result, begin_text = parse_docx_level_entry(filepath)
#    print(begin_text)
#    for i in result:
#        print(i)


















